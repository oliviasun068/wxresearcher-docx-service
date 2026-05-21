#!/usr/bin/env python3
from __future__ import annotations

import argparse
import base64
import json
import os
import re
import uuid
from http.server import BaseHTTPRequestHandler, ThreadingHTTPServer
from pathlib import Path
from typing import Any
from urllib.parse import unquote

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


ROOT = Path(__file__).resolve().parent
OUTPUT_DIR = ROOT / "out"
OUTPUT_DIR.mkdir(parents=True, exist_ok=True)


SECTION_ORDER = ["新锐观点", "市场动向", "成果发布"]

FIXED_EDITOR_NOTE = [
    "以习近平同志为核心的党中央高度重视建设世界一流企业。党的二十大报告明确提出，完善中国特色现代企业制度，弘扬企业家精神，加快建设世界一流企业。近年来，国务院国资委组织中央企业和地方国有重点企业开展对标世界一流管理提升行动，推动国有企业在先进的管理中要质量、要效益、要增长。",
    "作为服务党和国家科学民主决策的智力型中央企业，中咨公司始终以建设世界一流咨询机构和国家高端智库为目标，致力于探索如何建设一流咨询机构，如何在服务中国式现代化进程中实现再出发、再发展、再辉煌。为此，我们持续跟踪国内外20余家咨询机构（见附件）的新锐观点、市场动向及成果发布动态，深化对标研究，供交流参考。",
]

FIXED_TRACKING_SCOPE = [
    "麦肯锡（McKinsey & Company）",
    "贝恩公司（Bain & Company）",
    "波士顿咨询公司（BCG）",
    "埃森哲（Accenture）",
    "罗兰贝格管理咨询（Roland Berger）",
    "弗若斯特沙利文（Frost & Sullivan）",
    "德勤（Deloitte）",
    "IBM中国",
    "广东省国际工程咨询有限公司",
    "北京国际工程咨询有限公司",
    "国信国际工程咨询集团股份有限公司",
    "中国信息通信研究院",
    "北大纵横管理咨询集团",
    "综合开发研究院",
    "北京零点有数数据科技股份有限公司",
    "华夏基石管理咨询",
    "乔诺咨询",
    "和君咨询",
    "百思特管理咨询",
    "其他咨询机构",
]


def _clean_filename(value: str) -> str:
    value = re.sub(r"[\\/:*?\"<>|]+", "-", value).strip()
    return value[:80] or f"report-{uuid.uuid4().hex[:8]}"


def _as_str(value: Any) -> str:
    if value is None:
        return ""
    return str(value).strip()


def _as_list(value: Any) -> list[Any]:
    return value if isinstance(value, list) else []


def _parse_report_json(payload: Any) -> dict[str, Any]:
    if isinstance(payload, dict):
        return payload
    if isinstance(payload, str):
        text = payload.strip()
        if not text:
            raise ValueError("report_json is empty")
        try:
            parsed = json.loads(text)
            if isinstance(parsed, dict):
                return parsed
        except json.JSONDecodeError:
            pass
        match = re.search(r"\{[\s\S]*\}", text)
        if match:
            parsed = json.loads(match.group(0))
            if isinstance(parsed, dict):
                return parsed
    raise ValueError("report_json must be an object or a JSON object string")


def _normalize_report(raw: dict[str, Any]) -> dict[str, Any]:
    sections = []
    raw_sections = _as_list(raw.get("sections"))
    for sec in raw_sections:
        if not isinstance(sec, dict):
            continue
        title = _as_str(sec.get("section_title")).strip("【】")
        articles = []
        for article in _as_list(sec.get("articles")):
            if not isinstance(article, dict):
                continue
            detail = article.get("detail")
            if isinstance(detail, str):
                detail_parts = [p.strip() for p in re.split(r"\n{2,}", detail) if p.strip()]
            elif isinstance(detail, list):
                detail_parts = [_as_str(p) for p in detail if _as_str(p)]
            else:
                detail_parts = []

            item = {
                "title": _as_str(article.get("title")),
                "organization": _as_str(article.get("organization")),
                "publish_date": _as_str(article.get("publish_date")),
                "summary": _as_str(article.get("summary")),
                "detail": detail_parts,
                "url": _as_str(article.get("url")),
            }
            if item["title"] or item["summary"] or item["detail"]:
                articles.append(item)
        if title and articles:
            sections.append({"section_title": title, "articles": articles})

    ordered = []
    for name in SECTION_ORDER:
        ordered.extend([sec for sec in sections if sec["section_title"] == name])
    ordered.extend([sec for sec in sections if sec["section_title"] not in SECTION_ORDER])

    tracking_scope = []
    for item in _as_list(raw.get("tracking_scope")):
        text = _as_str(item)
        if text and text not in tracking_scope:
            tracking_scope.append(text)

    return {
        "title": _as_str(raw.get("title")),
        "editor_note": _as_str(raw.get("editor_note")),
        "sections": ordered,
        "tracking_scope": tracking_scope,
    }


def _clear_body(doc: Document) -> None:
    body = doc._element.body
    sect_pr = body.sectPr
    for child in list(body):
        if child is not sect_pr:
            body.remove(child)


def _new_document() -> Document:
    doc = Document()
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.46)
    section.bottom_margin = Cm(2.22)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(3.15)
    section.header_distance = Cm(0)
    section.footer_distance = Cm(1.75)
    return doc


def _set_run_font(
    run,
    size_pt: float | None = None,
    bold: bool | None = None,
    font: str = "仿宋",
    ascii_font: str | None = None,
) -> None:
    run.font.name = font
    if size_pt is not None:
        run.font.size = Pt(size_pt)
    if bold is not None:
        run.bold = bold
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.append(r_fonts)
    r_fonts.set(qn("w:eastAsia"), font)
    r_fonts.set(qn("w:ascii"), ascii_font or font)
    r_fonts.set(qn("w:hAnsi"), ascii_font or font)


def _format_para(
    paragraph,
    *,
    align=WD_ALIGN_PARAGRAPH.JUSTIFY,
    first_line_pt: float | None = 32,
    left_indent_pt: float | None = None,
    line_spacing: float | Pt = Pt(28),
    space_before: float = 0,
    space_after: float = 0,
) -> None:
    paragraph.alignment = align
    fmt = paragraph.paragraph_format
    fmt.line_spacing = line_spacing
    fmt.space_before = Pt(space_before)
    fmt.space_after = Pt(space_after)
    fmt.first_line_indent = Pt(first_line_pt) if first_line_pt is not None else None
    fmt.left_indent = Pt(left_indent_pt) if left_indent_pt is not None else None


def _add_paragraph(
    doc: Document,
    text: str = "",
    *,
    size=16,
    bold=False,
    font="仿宋",
    ascii_font: str | None = "Times New Roman",
    align=WD_ALIGN_PARAGRAPH.JUSTIFY,
    first_line_pt: float | None = 32,
    left_indent_pt: float | None = None,
    line_spacing: float | Pt = Pt(28),
    space_before: float = 0,
    space_after: float = 0,
):
    p = doc.add_paragraph()
    _format_para(
        p,
        align=align,
        first_line_pt=first_line_pt,
        left_indent_pt=left_indent_pt,
        line_spacing=line_spacing,
        space_before=space_before,
        space_after=space_after,
    )
    if text:
        run = p.add_run(text)
        _set_run_font(run, size, bold, font, ascii_font)
    return p


def _add_blank(doc: Document, *, line_spacing: float | Pt = Pt(28), align=WD_ALIGN_PARAGRAPH.JUSTIFY) -> None:
    _add_paragraph(doc, "", align=align, first_line_pt=None, line_spacing=line_spacing)


def _add_page_break(doc: Document) -> None:
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)


def _add_editor_note(doc: Document) -> None:
    for _ in range(5):
        _add_blank(doc)
    _add_paragraph(doc, "编者按", size=16, font="方正小标宋_GBK", first_line_pt=32)
    for part in FIXED_EDITOR_NOTE:
        _add_paragraph(doc, part, size=16, font="仿宋", first_line_pt=32)


def _estimate_article_pages(article: dict[str, Any]) -> int:
    text = " ".join(
        [
            article.get("title", ""),
            article.get("organization", ""),
            article.get("publish_date", ""),
            article.get("summary", ""),
            " ".join(article.get("detail", [])),
            article.get("url", ""),
        ]
    )
    return max(1, min(4, (len(text) + 650) // 900))


def _add_toc_article(doc: Document, title: str, page_no: int) -> None:
    p = _add_paragraph(
        doc,
        "",
        size=15.5,
        font="仿宋",
        align=WD_ALIGN_PARAGRAPH.LEFT,
        first_line_pt=None,
        line_spacing=1.8,
    )
    p.paragraph_format.tab_stops.add_tab_stop(
        Cm(15.3),
        alignment=WD_TAB_ALIGNMENT.RIGHT,
        leader=WD_TAB_LEADER.DOTS,
    )
    title_run = p.add_run(title)
    _set_run_font(title_run, 15.5, False, "仿宋", "Times New Roman")
    page_run = p.add_run(f"\t{page_no}")
    _set_run_font(page_run, 15.5, False, "仿宋", "Times New Roman")


def _add_toc(doc: Document, sections: list[dict[str, Any]]) -> None:
    _add_blank(doc, align=WD_ALIGN_PARAGRAPH.CENTER, line_spacing=1.0)
    _add_paragraph(
        doc,
        "目 录",
        size=18,
        bold=True,
        font="黑体",
        align=WD_ALIGN_PARAGRAPH.CENTER,
        first_line_pt=None,
        line_spacing=1.0,
    )
    _add_blank(doc, align=WD_ALIGN_PARAGRAPH.CENTER, line_spacing=1.0)
    page_no = 1
    for sec in sections:
        _add_paragraph(
            doc,
            f"【{sec['section_title']}】",
            size=15.5,
            bold=True,
            font="黑体",
            align=WD_ALIGN_PARAGRAPH.LEFT,
            first_line_pt=None,
            line_spacing=1.8,
        )
        for article in sec["articles"]:
            _add_toc_article(doc, article["title"], page_no)
            page_no += _estimate_article_pages(article)


def _add_section_heading(doc: Document, text: str) -> None:
    _add_blank(doc, line_spacing=Pt(28))
    _add_paragraph(
        doc,
        f"【{text}】",
        size=18,
        bold=True,
        font="微软雅黑",
        align=WD_ALIGN_PARAGRAPH.LEFT,
        first_line_pt=None,
        line_spacing=1.39,
    )


def _add_article_title(doc: Document, text: str) -> None:
    _add_paragraph(
        doc,
        text,
        size=18,
        bold=True,
        font="方正小标宋_GBK",
        align=WD_ALIGN_PARAGRAPH.CENTER,
        first_line_pt=None,
        line_spacing=Pt(28),
        space_before=15,
    )


def _add_meta(doc: Document, article: dict[str, Any]) -> None:
    parts = []
    if article["organization"]:
        parts.append(f"机构：{article['organization']}")
    if article["publish_date"]:
        parts.append(f"发布日期：{article['publish_date']}")
    if not parts:
        return
    _add_paragraph(
        doc,
        "|".join(parts),
        size=16,
        bold=True,
        font="楷体",
        align=WD_ALIGN_PARAGRAPH.CENTER,
        first_line_pt=None,
    )


def _add_summary(doc: Document, text: str) -> None:
    p = _add_paragraph(doc, "", size=16, font="楷体", first_line_pt=32)
    label = p.add_run("摘要")
    _set_run_font(label, 16, True, "黑体", "黑体")
    colon = p.add_run("：")
    _set_run_font(colon, 16, False, "黑体", "黑体")
    body = p.add_run(text)
    _set_run_font(body, 16, False, "楷体", "Times New Roman")


def _add_detail_label(doc: Document) -> None:
    _add_paragraph(
        doc,
        "详细内容",
        size=16,
        bold=True,
        font="黑体",
        align=WD_ALIGN_PARAGRAPH.LEFT,
        first_line_pt=32,
    )


def _add_body_text(doc: Document, text: str) -> None:
    for part in re.split(r"\n{2,}", _as_str(text)):
        if part.strip():
            _add_paragraph(doc, part.strip(), size=16, font="仿宋", first_line_pt=32)


def _add_tracking_scope(doc: Document) -> None:
    _add_paragraph(
        doc,
        "附件",
        size=15.5,
        font="黑体",
        align=WD_ALIGN_PARAGRAPH.JUSTIFY,
        first_line_pt=None,
        line_spacing=Pt(28),
    )
    p = _add_paragraph(
        doc,
        "主要跟踪范围",
        size=17.5,
        font="黑体",
        align=WD_ALIGN_PARAGRAPH.LEFT,
        first_line_pt=None,
        left_indent_pt=172.65,
        line_spacing=0.95,
        space_before=12.85,
    )
    p.paragraph_format.keep_with_next = True
    _add_blank(doc, line_spacing=Pt(12))
    for idx, name in enumerate(FIXED_TRACKING_SCOPE, 1):
        p = _add_paragraph(
            doc,
            "",
            size=15.5,
            font="仿宋",
            align=WD_ALIGN_PARAGRAPH.LEFT,
            first_line_pt=32,
            line_spacing=Pt(30),
        )
        p.paragraph_format.tab_stops.add_tab_stop(Cm(2.6), alignment=WD_TAB_ALIGNMENT.LEFT)
        num_run = p.add_run(f"{idx}.\t")
        _set_run_font(num_run, 15.5, False, "仿宋", "Times New Roman")
        name_run = p.add_run(name)
        _set_run_font(name_run, 15.5, False, "仿宋", "Times New Roman")


def render_docx(report: dict[str, Any]) -> Path:
    doc = _new_document()

    _add_editor_note(doc)
    _add_page_break(doc)
    _add_toc(doc, report["sections"])
    _add_page_break(doc)

    for sec in report["sections"]:
        _add_section_heading(doc, sec["section_title"])
        for article in sec["articles"]:
            _add_article_title(doc, article["title"])
            _add_meta(doc, article)
            if article["summary"]:
                _add_summary(doc, article["summary"])
            _add_blank(doc)
            _add_detail_label(doc)
            for para in article["detail"]:
                _add_body_text(doc, para)
            if article["url"]:
                _add_paragraph(doc, f"原文链接：{article['url']}", size=16, font="仿宋", first_line_pt=None)

    _add_page_break(doc)
    _add_tracking_scope(doc)

    filename = _clean_filename(report["title"] or "咨询行业动态追踪") + f"-{uuid.uuid4().hex[:8]}.docx"
    out_path = OUTPUT_DIR / filename
    doc.save(str(out_path))
    return out_path


class Handler(BaseHTTPRequestHandler):
    server_version = "WxResearcherDocxService/0.1"

    def _json(self, status: int, data: dict[str, Any]) -> None:
        raw = json.dumps(data, ensure_ascii=False).encode("utf-8")
        self.send_response(status)
        self.send_header("Content-Type", "application/json; charset=utf-8")
        self.send_header("Content-Length", str(len(raw)))
        self.end_headers()
        self.wfile.write(raw)

    def do_GET(self) -> None:
        if self.path == "/health":
            self._json(200, {"success": True, "message": "ok"})
            return
        if self.path.startswith("/files/"):
            name = unquote(self.path[len("/files/"):])
            path = (OUTPUT_DIR / name).resolve()
            if OUTPUT_DIR.resolve() not in path.parents or not path.exists():
                self.send_error(404)
                return
            data = path.read_bytes()
            self.send_response(200)
            self.send_header("Content-Type", "application/vnd.openxmlformats-officedocument.wordprocessingml.document")
            self.send_header("Content-Disposition", f'attachment; filename="{path.name}"')
            self.send_header("Content-Length", str(len(data)))
            self.end_headers()
            self.wfile.write(data)
            return
        self.send_error(404)

    def do_POST(self) -> None:
        if self.path != "/render":
            self.send_error(404)
            return
        try:
            length = int(self.headers.get("Content-Length", "0"))
            payload = json.loads(self.rfile.read(length).decode("utf-8"))
            raw_report = payload.get("report_json", payload)
            report = _normalize_report(_parse_report_json(raw_report))
            if not report["sections"]:
                raise ValueError("report_json.sections has no valid articles")
            out_path = render_docx(report)
            host = self.headers.get("Host", f"127.0.0.1:{self.server.server_port}")
            scheme = "https" if self.headers.get("X-Forwarded-Proto") == "https" else "http"
            file_url = f"{scheme}://{host}/files/{out_path.name}"
            encoded = base64.b64encode(out_path.read_bytes()).decode("ascii")
            self._json(200, {
                "success": True,
                "filename": out_path.name,
                "file_url": file_url,
                "docx_base64": encoded,
                "article_count": sum(len(sec["articles"]) for sec in report["sections"]),
                "section_count": len(report["sections"]),
            })
        except Exception as exc:
            self._json(400, {"success": False, "message": str(exc)})

    def log_message(self, fmt: str, *args: Any) -> None:
        print("%s - - [%s] %s" % (self.address_string(), self.log_date_time_string(), fmt % args))


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("--host", default=os.environ.get("HOST", "127.0.0.1"))
    parser.add_argument("--port", type=int, default=int(os.environ.get("PORT", "8765")))
    args = parser.parse_args()

    server = ThreadingHTTPServer((args.host, args.port), Handler)
    print(f"Serving DOCX service on http://{args.host}:{args.port}")
    print("Template: generated in code")
    print(f"Output: {OUTPUT_DIR}")
    server.serve_forever()


if __name__ == "__main__":
    main()
